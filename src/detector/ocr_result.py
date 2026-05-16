import json
import re
from pathlib import Path
from paddleocr import PaddleOCR
from PIL import Image
import cv2
import numpy as np
from tqdm import tqdm
import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from config.config import OUTPUT_DIR
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Khởi tạo OCR cho cả Việt và Anh
ocr_vi = PaddleOCR(
    use_textline_orientation=True,
    lang='vi',
)

ocr_en = PaddleOCR(
    use_textline_orientation=True,
    lang='en',
)

# Chiều cao tối thiểu để PaddleOCR recognizer nhìn rõ chữ.
# Cell bảng kỹ thuật thường rất nhỏ; nếu không upscale dễ mất ký tự.
MIN_OCR_HEIGHT = 48

# Regex bắt ký tự alpha-numeric + chữ tiếng Việt có dấu, dùng để
# đếm số ký tự "có nghĩa" khi cho điểm kết quả OCR.
_ALNUM_VN_RE = re.compile(r"[\w\u00C0-\u1EF9]")


def _preprocess_for_ocr(image_path):
    """Đọc ảnh cell, upscale nếu quá nhỏ và khử nhiễu nhẹ trước khi OCR.

    Trả về numpy BGR để PaddleOCR có thể nhận trực tiếp.
    """
    img = cv2.imread(str(image_path))
    if img is None:
        return None

    h, w = img.shape[:2]
    if h > 0 and h < MIN_OCR_HEIGHT:
        scale = MIN_OCR_HEIGHT / float(h)
        new_w = max(1, int(round(w * scale)))
        img = cv2.resize(img, (new_w, MIN_OCR_HEIGHT), interpolation=cv2.INTER_CUBIC)

    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    gray = cv2.bilateralFilter(gray, 5, 50, 50)
    return cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)


def _score_lines(lines):
    """Cho điểm tập kết quả OCR dựa trên độ tự tin trung bình và số ký tự thật.

    Việc kết hợp mean_conf với sqrt(char_count) giúp:
    - Không bị thiên vị model trả ra nhiều cụm nhiễu nhỏ (như khi dùng sum).
    - Vẫn ưu tiên model đọc được nhiều ký tự ý nghĩa.
    """
    if not lines:
        return 0.0
    confs = [c for _, _, c in lines]
    mean_conf = sum(confs) / len(confs)
    char_count = sum(len(_ALNUM_VN_RE.findall(t)) for _, t, _ in lines)
    return mean_conf * (char_count ** 0.5)


def _collect_lines(result):
    """Chuẩn hoá output PaddleOCR thành list (y_center, text, conf)."""
    lines = []
    if not result or not result[0]:
        return lines
    for line in result[0]:
        text = line[1][0]
        conf = line[1][1]
        y_center = (line[0][0][1] + line[0][2][1]) / 2
        lines.append((y_center, text, conf))
    return lines


def ocr_cell(image_path):
    """OCR cho một ảnh cell, trả về toàn bộ text (nối các dòng) và confidence trung bình"""
    try:
        prepared = _preprocess_for_ocr(image_path)
        ocr_input = prepared if prepared is not None else str(image_path)

        result_vi = ocr_vi.ocr(ocr_input)
        result_en = ocr_en.ocr(ocr_input)

        lines_vi = _collect_lines(result_vi)
        lines_en = _collect_lines(result_en)

        # Chọn model nào có điểm tổng hợp cao hơn (mean_conf * sqrt(char_count)).
        score_vi = _score_lines(lines_vi)
        score_en = _score_lines(lines_en)
        chosen_lines = lines_vi if score_vi >= score_en else lines_en

        if not chosen_lines:
            return "", 0

        # Sắp xếp theo vị trí y (từ trên xuống dưới)
        chosen_lines.sort(key=lambda x: x[0])

        # Nối tất cả các dòng
        full_text = "\n".join(text for _, text, _ in chosen_lines)
        avg_conf = sum(c for _, _, c in chosen_lines) / len(chosen_lines)

        return full_text, avg_conf
    except Exception as e:
        print(f"Loi OCR {image_path}: {e}")
        return "", 0


# Các class chứa text cần OCR (bao gồm cả tên gốc và tên đã mapping)
OCR_CLASSES = {
    'Text', 'List-item', 'Note',
    'Section-header', 'Page-header', 'Page-footer',
    'Title', 'Caption', 'Formula'
}

# Các class KHÔNG cần OCR (ảnh, bảng sẽ xử lý riêng ở Model 2)
SKIP_OCR_CLASSES = {'PartDrawing', 'Picture', 'Table'}


def process_all_jsons():
    # Đường dẫn lấy từ config — không hardcode
    path_out1 = OUTPUT_DIR
    path_cropped_model2 = OUTPUT_DIR / "cropped_model2"

    # Tìm tất cả JSON files
    json_files_model1 = list(path_out1.glob("*.json"))
    json_files_model2 = list(path_cropped_model2.glob("*_metadata.json")) if path_cropped_model2.exists() else []

    print(f"Tim thay {len(json_files_model1)} JSON tu Model 1")
    print(f"Tim thay {len(json_files_model2)} JSON tu Model 2")

    # Kết quả tổng hợp
    final_result = {
        "metadata": {
            "total_model1_jsons": len(json_files_model1),
            "total_model2_jsons": len(json_files_model2),
            "ocr_languages": ["vi", "en"]
        },
        "model1_results": [],
        "model2_results": []
    }

    # XỬ LÝ MODEL 1 - OCR tất cả vùng có chữ
    print("\n=== XU LY MODEL 1 (Tat ca vung co chu) ===")
    for json_path in tqdm(json_files_model1, desc="Model 1"):
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        processed_objects = []

        for obj in data.get('objects', []):
            obj_class = obj.get('class', '')
            original_class = obj.get('original_class', obj_class)

            # OCR tất cả class có chứa text (dùng cả class đã map và class gốc)
            if obj_class in OCR_CLASSES or original_class in OCR_CLASSES:
                crop_path = obj.get('crop_path')
                if crop_path and Path(crop_path).exists():
                    text, confidence = ocr_cell(crop_path)
                    obj['ocr_text'] = text
                    obj['ocr_confidence'] = confidence
                    preview = text[:60].replace('\n', ' ')
                    print(f"  [{obj_class}] '{preview}...' (conf: {confidence:.2f})")
                else:
                    obj['ocr_text'] = ""
                    obj['ocr_confidence'] = 0
                    print(f"  Khong tim thay anh: {crop_path}")
            elif obj_class not in SKIP_OCR_CLASSES:
                # Class không xác định — vẫn thử OCR
                crop_path = obj.get('crop_path')
                if crop_path and Path(crop_path).exists():
                    text, confidence = ocr_cell(crop_path)
                    obj['ocr_text'] = text
                    obj['ocr_confidence'] = confidence

            processed_objects.append(obj)

        data['objects'] = processed_objects
        final_result['model1_results'].append(data)

    # XỬ LÝ MODEL 2 - Xử lý tất cả cells
    print("\n=== XU LY MODEL 2 (Tat ca cells) ===")
    for json_path in tqdm(json_files_model2, desc="Model 2"):
        with open(json_path, 'r', encoding='utf-8') as f:
            cells_data = json.load(f)

        # cells_data có thể là list hoặc dict
        if isinstance(cells_data, dict):
            cells = cells_data.get('cells', [])
        else:
            cells = cells_data

        processed_cells = []

        for cell in cells:
            crop_path = cell.get('cropped_image_path')
            if crop_path and Path(crop_path).exists():
                text, confidence = ocr_cell(crop_path)
                cell['ocr_text'] = text
                cell['ocr_confidence'] = confidence
                print(f"  Cell {cell.get('cell_index', '?')}: '{text}' (conf: {confidence:.2f})")
            else:
                cell['ocr_text'] = ""
                cell['ocr_confidence'] = 0

            processed_cells.append(cell)

        if isinstance(cells_data, dict):
            cells_data['cells'] = processed_cells
            final_result['model2_results'].append(cells_data)
        else:
            final_result['model2_results'].append(processed_cells)

    # LƯU KẾT QUẢ
    output_dir = OUTPUT_DIR / "ocr_results"
    output_dir.mkdir(parents=True, exist_ok=True)

    # Lưu file JSON tổng hợp
    final_json_path = output_dir / "all_ocr_results.json"
    with open(final_json_path, 'w', encoding='utf-8') as f:
        json.dump(final_result, f, ensure_ascii=False, indent=2)

    print(f"\nDa luu ket qua tong hop tai: {final_json_path}")

    # Tạo file JSON riêng cho web (đơn giản hơn)
    web_json_path = output_dir / "web_demo_results.json"
    web_data = {
        "total_ocr_cells": 0,
        "results": []
    }

    # Gom tất cả kết quả OCR từ model 1
    for data in final_result['model1_results']:
        for obj in data.get('objects', []):
            cls = obj.get('class', '')
            orig_cls = obj.get('original_class', '')
            
            # Text objects
            if 'ocr_text' in obj and obj['ocr_text']:
                web_data['results'].append({
                    "source": "model1",
                    "type": "text",
                    "class": cls,
                    "bbox": obj.get('bbox'),
                    "text": obj['ocr_text'],
                    "confidence": obj.get('ocr_confidence', 0),
                    "image_path": obj.get('crop_path')
                })
            # Picture objects
            elif cls in ['Picture', 'PartDrawing'] or orig_cls in ['Picture', 'PartDrawing']:
                web_data['results'].append({
                    "source": "model1",
                    "type": "image",
                    "class": cls,
                    "bbox": obj.get('bbox'),
                    "text": "[HÌNH ẢNH]",
                    "confidence": obj.get('confidence', 0),
                    "image_path": obj.get('crop_path')
                })

    # Gom tất cả kết quả OCR từ model 2
    for cells in final_result['model2_results']:
        if isinstance(cells, dict):
            cells_list = cells.get('cells', [])
        else:
            cells_list = cells

        for cell in cells_list:
            if 'ocr_text' in cell and cell['ocr_text']:
                web_data['results'].append({
                    "source": "model2",
                    "type": "text",
                    "class": "TableCell",
                    "bbox": cell.get('bbox'),
                    "text": cell['ocr_text'],
                    "confidence": cell.get('ocr_confidence', 0),
                    "image_path": cell.get('cropped_image_path')
                })

    web_data['total_ocr_cells'] = len(web_data['results'])

    with open(web_json_path, 'w', encoding='utf-8') as f:
        json.dump(web_data, f, ensure_ascii=False, indent=2)

    print(f"Da luu file cho web demo tai: {web_json_path}")
    print(f"Tong so cells co OCR: {web_data['total_ocr_cells']}")

    # === XUẤT FILE TEXT GIỮ CẤU TRÚC ===
    export_structured_text(final_result, output_dir)

    # === TÁI TẠO CẤU TRÚC DOCUMENT BẰNG ENGINE MỚI ===
    print("\n=== KHOI DONG DOCUMENT RECONSTRUCTION ENGINE ===")
    from src.reconstruction.engine import DocumentReconstructor
    engine = DocumentReconstructor(output_dir)
    engine.process(final_result)

    # === XUẤT FILE DOCX THEO CÁCH CŨ (Optional) ===
    # export_docx(final_result, output_dir)

    return final_result


def export_structured_text(final_result, output_dir):
    """
    Xuất kết quả OCR thành file .txt giữ nguyên cấu trúc layout
    dựa trên tọa độ bounding box (sắp xếp từ trên xuống dưới, trái sang phải).
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    for img_data in final_result.get('model1_results', []):
        image_name = Path(img_data.get('image', 'unknown')).stem
        image_width = img_data.get('image_info', {}).get('width', 1000)
        image_height = img_data.get('image_info', {}).get('height', 1000)

        # Thu thập tất cả text blocks có OCR
        text_blocks = []
        for obj in img_data.get('objects', []):
            ocr_text = obj.get('ocr_text', '').strip()
            if not ocr_text:
                continue

            bbox = obj.get('bbox', {})
            text_blocks.append({
                'class': obj.get('class', ''),
                'original_class': obj.get('original_class', ''),
                'text': ocr_text,
                'confidence': obj.get('ocr_confidence', 0),
                'x1': bbox.get('x1', 0),
                'y1': bbox.get('y1', 0),
                'x2': bbox.get('x2', 0),
                'y2': bbox.get('y2', 0),
            })

        if not text_blocks:
            print(f"  [TEXT] {image_name}: khong co text nao de xuat")
            continue

        # ====================================================
        # Thuật toán sắp xếp giữ cấu trúc layout:
        # 1. Gom các block có y1 gần nhau thành cùng 1 "dòng"
        # 2. Trong mỗi dòng, sắp xếp theo x1 (trái -> phải)
        # 3. Giữa các dòng, sắp xếp theo y1 (trên -> dưới)
        # ====================================================

        # Sắp xếp theo y1 trước
        text_blocks.sort(key=lambda b: b['y1'])

        # Gom thành các dòng (row) — các block có y1 chênh nhau < threshold
        # thì coi như cùng 1 dòng
        LINE_THRESHOLD = image_height * 0.02  # 2% chiều cao ảnh
        rows = []
        current_row = [text_blocks[0]]

        for block in text_blocks[1:]:
            if abs(block['y1'] - current_row[0]['y1']) < LINE_THRESHOLD:
                current_row.append(block)
            else:
                rows.append(current_row)
                current_row = [block]
        rows.append(current_row)

        # Trong mỗi dòng, sắp xếp theo x1
        for row in rows:
            row.sort(key=lambda b: b['x1'])

        # ====================================================
        # Xuất file .txt
        # ====================================================
        txt_path = output_dir / f"{image_name}_structured.txt"
        lines_output = []

        lines_output.append(f"{'=' * 70}")
        lines_output.append(f"  KET QUA OCR - {img_data.get('image', '')}")
        lines_output.append(f"  Kich thuoc anh: {image_width} x {image_height}")
        lines_output.append(f"  So vung phat hien: {img_data.get('num_objects', 0)}")
        lines_output.append(f"  So vung co text: {len(text_blocks)}")
        lines_output.append(f"{'=' * 70}")
        lines_output.append("")

        prev_row_y = 0
        for row_idx, row in enumerate(rows):
            # Thêm khoảng trống nếu có khoảng cách lớn giữa các dòng
            row_y = row[0]['y1']
            gap = row_y - prev_row_y
            if prev_row_y > 0 and gap > image_height * 0.05:
                lines_output.append("")  # dòng trống = khoảng cách lớn

            # Xử lý từng block trong dòng
            row_texts = []
            for block in row:
                cls = block['class']
                text = block['text']

                # Format theo loại
                if cls in ('Page-header', 'Title'):
                    text = f"[HEADER] {text}"
                elif cls in ('Page-footer',):
                    text = f"[FOOTER] {text}"
                elif cls in ('Section-header',):
                    text = f"[SECTION] {text}"
                elif cls in ('Table',):
                    text = f"[TABLE]\n{text}"

                row_texts.append(text)

            # Nối các text trong cùng 1 dòng bằng dấu tab
            line = "    ".join(row_texts)
            lines_output.append(line)
            prev_row_y = row_y

        lines_output.append("")
        lines_output.append(f"{'=' * 70}")
        lines_output.append(f"  KET THUC - {image_name}")
        lines_output.append(f"{'=' * 70}")

        # Ghi file
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines_output))

        print(f"  Exported text file: {txt_path}")

    # ====================================================
    # File tổng hợp tất cả ảnh
    # ====================================================
    combined_path = output_dir / "ket_qua_tong_hop.txt"
    all_lines = []
    all_lines.append("KET QUA OCR TONG HOP - Engineering Drawings")
    all_lines.append(f"So anh xu ly: {len(final_result.get('model1_results', []))}")
    all_lines.append(f"{'=' * 70}")
    all_lines.append("")

    for img_data in final_result.get('model1_results', []):
        image_name = Path(img_data.get('image', 'unknown')).stem
        txt_path = output_dir / f"{image_name}_structured.txt"
        if txt_path.exists():
            with open(txt_path, 'r', encoding='utf-8') as f:
                all_lines.append(f.read())
            all_lines.append("\n")

    with open(combined_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(all_lines))

    print(f"\n  Exported combined file: {combined_path}")


def export_docx(final_result, output_dir):
    """
    Xuất kết quả OCR thành file .docx giữ nguyên cấu trúc layout.
    - Page-header -> Heading 1
    - Section-header -> Heading 2
    - Note/Text -> Normal paragraph
    - Page-footer -> Small italic
    - Table -> Word Table
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    # ========== FILE DOCX CHO TỪNG ẢNH ==========
    for img_data in final_result.get('model1_results', []):
        image_name = Path(img_data.get('image', 'unknown')).stem
        image_width = img_data.get('image_info', {}).get('width', 1000)
        image_height = img_data.get('image_info', {}).get('height', 1000)

        # Thu thập tất cả text blocks có OCR và cả ảnh (Picture)
        text_blocks = []
        for obj in img_data.get('objects', []):
            ocr_text = obj.get('ocr_text', '').strip()
            cls = obj.get('class', '')
            orig_cls = obj.get('original_class', '')
            crop_path = obj.get('crop_path', '')

            # Bỏ qua nếu không có text VÀ không phải là hình ảnh
            is_image = cls in ['Picture', 'PartDrawing'] or orig_cls in ['Picture', 'PartDrawing']
            if not ocr_text and not is_image:
                continue

            bbox = obj.get('bbox', {})
            text_blocks.append({
                'class': 'Picture' if is_image else cls,
                'original_class': orig_cls,
                'text': ocr_text,
                'crop_path': crop_path,
                'confidence': obj.get('ocr_confidence', 0),
                'x1': bbox.get('x1', 0),
                'y1': bbox.get('y1', 0),
                'x2': bbox.get('x2', 0),
                'y2': bbox.get('y2', 0),
            })

        if not text_blocks:
            continue

        # Sắp xếp theo y1, rồi theo x1
        text_blocks.sort(key=lambda b: (b['y1'], b['x1']))

        # Gom thành các dòng (row) — y1 chênh < threshold = cùng dòng
        LINE_THRESHOLD = image_height * 0.02
        rows = []
        current_row = [text_blocks[0]]

        for block in text_blocks[1:]:
            if abs(block['y1'] - current_row[0]['y1']) < LINE_THRESHOLD:
                current_row.append(block)
            else:
                rows.append(current_row)
                current_row = [block]
        rows.append(current_row)

        # Trong mỗi dòng, sắp xếp theo x1
        for row in rows:
            row.sort(key=lambda b: b['x1'])

        # ========== TẠO DOCUMENT ==========
        doc = Document()

        # Cài đặt font mặc định
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Tiêu đề tài liệu
        title = doc.add_heading(f'OCR - {img_data.get("image", "")}', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Thông tin ảnh
        info_para = doc.add_paragraph()
        info_run = info_para.add_run(
            f'Kich thuoc: {image_width} x {image_height} | '
            f'So vung: {img_data.get("num_objects", 0)} | '
            f'Co text: {len(text_blocks)}'
        )
        info_run.font.size = Pt(9)
        info_run.font.color.rgb = RGBColor(128, 128, 128)
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # dòng trống

        # ========== XUẤT TỪNG DÒNG ==========
        prev_row_y = 0
        for row in rows:
            row_y = row[0]['y1']
            gap = row_y - prev_row_y

            # Khoảng trống lớn giữa các vùng -> thêm dòng trống
            if prev_row_y > 0 and gap > image_height * 0.06:
                doc.add_paragraph()

            for block in row:
                cls = block['class']
                text = block['text']

                if cls in ('Page-header', 'Title'):
                    # -> Heading 1
                    h = doc.add_heading(text, level=1)
                    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

                elif cls in ('Section-header',):
                    # -> Heading 2
                    h = doc.add_heading(text, level=2)
                    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

                elif cls in ('Page-footer',):
                    # -> Nhỏ, nghiêng, cuối trang
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(100, 100, 100)

                elif cls in ('Table',):
                    # Nội dung bảng: thử tách thành các dòng và cột
                    lines = text.split('\n')
                    if len(lines) > 1:
                        # Thử detect cấu trúc bảng bằng tab/nhiều space
                        table_rows = []
                        for line in lines:
                            # Tách theo tab hoặc 2+ spaces
                            import re
                            cells = re.split(r'\t|  +', line.strip())
                            cells = [c.strip() for c in cells if c.strip()]
                            if cells:
                                table_rows.append(cells)

                        if table_rows:
                            max_cols = max(len(r) for r in table_rows)
                            # Pad các dòng ngắn
                            for r in table_rows:
                                while len(r) < max_cols:
                                    r.append('')

                            # Tạo bảng Word
                            table = doc.add_table(
                                rows=len(table_rows),
                                cols=max_cols
                            )
                            table.style = 'Table Grid'
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER

                            for i, row_data in enumerate(table_rows):
                                for j, cell_text in enumerate(row_data):
                                    cell = table.cell(i, j)
                                    cell.text = cell_text
                                    # Font cho cell
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.size = Pt(10)
                                            run.font.name = 'Times New Roman'
                        else:
                            p = doc.add_paragraph(text)
                    else:
                        p = doc.add_paragraph(text)

                elif cls in ('Picture', 'PartDrawing'):
                    # Chèn ảnh trực tiếp vào file Word
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    crop = block.get('crop_path')
                    if crop and Path(crop).exists():
                        try:
                            # Đặt kích thước ảnh vừa phải trong Word (tối đa ~5-6 inch)
                            p.add_run().add_picture(str(crop), width=Inches(5.5))
                        except Exception as e:
                            p.add_run(f"[LỖI CHÈN ẢNH: {e}]")
                    else:
                        p.add_run("[KHÔNG TÌM THẤY ẢNH CROP]")

                else:
                    # Note, Text, hoặc class khác -> Normal paragraph
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'

            prev_row_y = row_y

        # Lưu file
        docx_path = output_dir / f"{image_name}.docx"
        doc.save(str(docx_path))
        print(f"  Exported DOCX: {docx_path}")

    # ========== FILE DOCX TỔNG HỢP ==========
    doc_all = Document()

    style = doc_all.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    doc_all.add_heading('KET QUA OCR TONG HOP', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = doc_all.add_paragraph()
    info.add_run(
        f'So anh xu ly: {len(final_result.get("model1_results", []))}'
    ).font.size = Pt(10)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_all.add_paragraph()

    for img_data in final_result.get('model1_results', []):
        image_name = Path(img_data.get('image', 'unknown')).stem
        single_docx = output_dir / f"{image_name}.docx"

        if single_docx.exists():
            # Đọc lại file đơn và nối vào
            sub_doc = Document(str(single_docx))
            for element in sub_doc.element.body:
                doc_all.element.body.append(element)

            # Thêm page break giữa các ảnh
            doc_all.add_page_break()

    combined_docx = output_dir / "ket_qua_tong_hop.docx"
    doc_all.save(str(combined_docx))
    print(f"  Exported combined DOCX: {combined_docx}")


# Chạy xử lý
if __name__ == "__main__":
    result = process_all_jsons()