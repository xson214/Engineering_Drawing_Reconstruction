import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

class DocumentReconstructor:
    def __init__(self, output_dir):
        self.output_dir = Path(output_dir)

    def process(self, final_result):
        """
        Xử lý final_result (chứa model1_results và model2_results)
        và trả về danh sách các document đã reconstruct.
        """
        # Nhóm model2 cells theo image_stem và table_id
        # model2_results là list các dict hoặc list các list
        # format: [{"cells": [{"cropped_image_path": ".../2_table_1_model2_crop_...jpg", ...}]}]
        # hoặc [[{"cropped_image_path": "..."}]]
        model2_grouped = {}
        for item in final_result.get("model2_results", []):
            if isinstance(item, dict):
                cells = item.get("cells", [])
                rows_meta = item.get("rows", [])
                cols_meta = item.get("cols", [])
            else:
                cells = item
                rows_meta, cols_meta = [], []
            if not cells:
                continue

            # Lấy path của cell đầu tiên để trích xuất thông tin
            crop_path = Path(cells[0].get("cropped_image_path", ""))
            if not crop_path.name:
                continue

            # Extract image_name and table_id. Pattern: {image_name}_table_{i}_...
            match = re.match(r"(.*)_table_(\d+)_model2_crop", crop_path.name)
            if match:
                img_stem = match.group(1)
                table_idx = int(match.group(2))

                if img_stem not in model2_grouped:
                    model2_grouped[img_stem] = {}
                model2_grouped[img_stem][table_idx] = {
                    "cells": cells,
                    "rows": rows_meta,
                    "cols": cols_meta,
                }

        reconstructed_docs = []
        for img_data in final_result.get("model1_results", []):
            doc_struct = self._reconstruct_single_image(img_data, model2_grouped)
            reconstructed_docs.append(doc_struct)
            
            # Export JSON
            json_path = self.output_dir / f"{Path(doc_struct['document']['image_name']).stem}_structured.json"
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(doc_struct, f, ensure_ascii=False, indent=2)
                
            # Export DOCX
            self._export_to_docx(doc_struct)
            
        return reconstructed_docs

    def _reconstruct_single_image(self, img_data, model2_grouped):
        image_name = img_data.get('image', 'unknown')
        img_stem = Path(image_name).stem
        width = img_data.get('image_info', {}).get('width', 1000)
        height = img_data.get('image_info', {}).get('height', 1000)
        
        doc_struct = {
            "document": {
                "image_name": image_name,
                "width": width,
                "height": height
            },
            "elements": []
        }
        
        objects = img_data.get('objects', [])
        
        tables_raw = []
        notes_raw = []
        drawings_raw = []
        others_raw = []
        
        for obj in objects:
            cls = obj.get('class', '')
            orig_cls = obj.get('original_class', '')
            
            if cls == 'Table' or orig_cls == 'Table':
                tables_raw.append(obj)
            elif cls in ['List-item', 'Text', 'Note'] or orig_cls in ['List-item', 'Text', 'Note']:
                notes_raw.append(obj)
            elif cls in ['Picture', 'PartDrawing'] or orig_cls in ['Picture', 'PartDrawing']:
                drawings_raw.append(obj)
            else:
                others_raw.append(obj)
                
        # 1. Xử lý Tables
        # Sort objects top-to-bottom for predictable mapping
        tables_raw.sort(key=lambda x: (x.get('bbox', {}).get('y1', 0), x.get('bbox', {}).get('x1', 0)))
        
        tables_cells_dict = model2_grouped.get(img_stem, {})
        
        for i, table_obj in enumerate(tables_raw):
            table_idx = i + 1
            table_data = tables_cells_dict.get(table_idx,
                                               {"cells": [], "rows": [], "cols": []})
            # Tương thích format cũ (chỉ là list cells).
            if isinstance(table_data, list):
                table_data = {"cells": table_data, "rows": [], "cols": []}

            bbox = table_obj.get('bbox', {'x1': 0, 'y1': 0, 'x2': 0, 'y2': 0})

            structure = self._reconstruct_table_grid(table_data, bbox)
            
            norm_bbox = [
                bbox['x1'] / width, bbox['y1'] / height,
                bbox['x2'] / width, bbox['y2'] / height
            ]
            
            doc_struct["elements"].append({
                "id": f"table_{table_obj.get('id', i)}",
                "type": "table",
                "bbox": [bbox['x1'], bbox['y1'], bbox['x2'], bbox['y2']],
                "normalized_bbox": [round(n, 4) for n in norm_bbox],
                "structure": structure,
                "confidence": table_obj.get('confidence', 0)
            })
            
        # 2. Xử lý Notes
        # Gom cụm các notes (List-item, Text) gần nhau theo chiều dọc
        notes_raw.sort(key=lambda x: x.get('bbox', {}).get('y1', 0))
        
        grouped_notes = []
        current_group = []
        
        for note in notes_raw:
            if not current_group:
                current_group.append(note)
                continue
                
            last_note = current_group[-1]
            last_bbox = last_note.get('bbox', {})
            note_bbox = note.get('bbox', {})
            
            x_overlap = max(0, min(last_bbox.get('x2', 0), note_bbox.get('x2', 0)) - max(last_bbox.get('x1', 0), note_bbox.get('x1', 0)))
            x_range_min = min(last_bbox.get('x2', 0) - last_bbox.get('x1', 0), note_bbox.get('x2', 0) - note_bbox.get('x1', 0))
            
            y_gap = note_bbox.get('y1', 0) - last_bbox.get('y2', 0)
            
            if x_range_min > 0 and (x_overlap / x_range_min) > 0.3 and y_gap < height * 0.05:
                current_group.append(note)
            else:
                grouped_notes.append(current_group)
                current_group = [note]
                
        if current_group:
            grouped_notes.append(current_group)
            
        for i, group in enumerate(grouped_notes):
            x1 = min(n.get('bbox', {}).get('x1', 0) for n in group)
            y1 = min(n.get('bbox', {}).get('y1', 0) for n in group)
            x2 = max(n.get('bbox', {}).get('x2', 0) for n in group)
            y2 = max(n.get('bbox', {}).get('y2', 0) for n in group)
            
            items = [n.get('ocr_text', '') for n in group if n.get('ocr_text', '').strip()]
            
            if not items:
                continue
                
            norm_bbox = [x1 / width, y1 / height, x2 / width, y2 / height]
            doc_struct["elements"].append({
                "id": f"notes_{i}",
                "type": "notes",
                "bbox": [x1, y1, x2, y2],
                "normalized_bbox": [round(n, 4) for n in norm_bbox],
                "items": items
            })
            
        # 3. Xử lý Drawing
        for i, draw_obj in enumerate(drawings_raw):
            bbox = draw_obj.get('bbox', {'x1': 0, 'y1': 0, 'x2': 0, 'y2': 0})
            norm_bbox = [
                bbox['x1'] / width, bbox['y1'] / height,
                bbox['x2'] / width, bbox['y2'] / height
            ]
            doc_struct["elements"].append({
                "id": f"drawing_{draw_obj.get('id', i)}",
                "type": "drawing",
                "bbox": [bbox['x1'], bbox['y1'], bbox['x2'], bbox['y2']],
                "normalized_bbox": [round(n, 4) for n in norm_bbox],
                "image_path": draw_obj.get('crop_path', '')
            })
            
        # 4. Xử lý Others (Headers, Metadata)
        for i, other in enumerate(others_raw):
            bbox = other.get('bbox', {'x1': 0, 'y1': 0, 'x2': 0, 'y2': 0})
            norm_bbox = [
                bbox['x1'] / width, bbox['y1'] / height,
                bbox['x2'] / width, bbox['y2'] / height
            ]
            doc_struct["elements"].append({
                "id": f"other_{other.get('id', i)}",
                "type": "metadata",
                "class_name": other.get('class', ''),
                "bbox": [bbox['x1'], bbox['y1'], bbox['x2'], bbox['y2']],
                "normalized_bbox": [round(n, 4) for n in norm_bbox],
                "text": other.get('ocr_text', '')
            })
            
        # Sắp xếp reading order: top-to-bottom, left-to-right
        doc_struct["elements"].sort(key=lambda e: (e['bbox'][1], e['bbox'][0]))
        
        return doc_struct

    def _reconstruct_table_grid(self, table_data, table_bbox):
        """Dựng cấu trúc lưới bảng từ metadata Model 2.

        Nếu có row/col bbox từ TableDetector → dùng trực tiếp (chính xác cả
        khi có cột rỗng hoặc cell rỗng). Ngược lại fallback về clustering
        x_center theo cell có chữ như logic cũ.
        """
        if isinstance(table_data, dict):
            cells = table_data.get('cells', [])
            rows_meta = table_data.get('rows', [])
            cols_meta = table_data.get('cols', [])
        else:
            cells = table_data
            rows_meta, cols_meta = [], []

        if not cells:
            return {"columns": 0, "rows": [], "col_widths": [], "row_heights": []}

        if rows_meta and cols_meta:
            return self._build_grid_from_lines(cells, rows_meta, cols_meta)
        return self._build_grid_from_cells(cells, table_bbox)

    def _build_grid_from_lines(self, cells, rows_meta, cols_meta):
        """Dựng lưới bảng dựa trực tiếp vào bbox row/col của detector.

        Đây là path chính: số cột/độ rộng cột phản ánh đúng cấu trúc bảng gốc
        kể cả khi có cột rỗng (vd. cột GRADE) hoặc cell rỗng.
        """
        rows_sorted = sorted(self._merge_lines(rows_meta, axis='y', tol=8),
                             key=lambda r: r['bbox'][1])
        cols_sorted = sorted(self._merge_lines(cols_meta, axis='x', tol=8),
                             key=lambda c: c['bbox'][0])

        num_rows = len(rows_sorted)
        num_cols = len(cols_sorted)
        if num_rows == 0 or num_cols == 0:
            return {"columns": 0, "rows": [], "col_widths": [], "row_heights": []}

        col_widths = [max(1, c['bbox'][2] - c['bbox'][0]) for c in cols_sorted]
        row_heights = [max(1, r['bbox'][3] - r['bbox'][1]) for r in rows_sorted]

        row_ranges = [(r['bbox'][1], r['bbox'][3]) for r in rows_sorted]
        col_ranges = [(c['bbox'][0], c['bbox'][2]) for c in cols_sorted]

        grid = [["" for _ in range(num_cols)] for _ in range(num_rows)]
        for cell in cells:
            bbox = cell.get('bbox') or []
            if len(bbox) < 4:
                continue
            text = (cell.get('ocr_text') or '').strip()
            if not text:
                continue
            cx = (bbox[0] + bbox[2]) / 2
            cy = (bbox[1] + bbox[3]) / 2
            row_idx = self._best_index(cy, row_ranges)
            col_idx = self._best_index(cx, col_ranges)
            if row_idx is None or col_idx is None:
                continue
            existing = grid[row_idx][col_idx]
            grid[row_idx][col_idx] = (existing + ' ' + text).strip() if existing else text

        return {
            "columns": num_cols,
            "rows": grid,
            "col_widths": col_widths,
            "row_heights": row_heights,
        }

    @staticmethod
    def _merge_lines(lines, axis='y', tol=8):
        """Gộp các bbox row/col gần trùng (tránh sinh hàng/cột dư)."""
        if not lines:
            return []
        key_lo = 1 if axis == 'y' else 0
        key_hi = 3 if axis == 'y' else 2
        sorted_lines = sorted(lines, key=lambda d: d['bbox'][key_lo])
        merged = []
        for ln in sorted_lines:
            bbox = list(ln['bbox'])
            conf = float(ln.get('confidence', 0))
            if merged and abs(bbox[key_lo] - merged[-1]['bbox'][key_lo]) <= tol:
                last = merged[-1]
                last['bbox'][key_lo] = min(last['bbox'][key_lo], bbox[key_lo])
                last['bbox'][key_hi] = max(last['bbox'][key_hi], bbox[key_hi])
                last['confidence'] = max(last['confidence'], conf)
                continue
            merged.append({"bbox": bbox, "confidence": conf})
        return merged

    @staticmethod
    def _best_index(value, ranges):
        """Tìm chỉ số range chứa `value`; nếu không có thì chọn range gần nhất."""
        if not ranges:
            return None
        for i, (lo, hi) in enumerate(ranges):
            if lo <= value <= hi:
                return i
        best_idx = 0
        best_dist = float('inf')
        for i, (lo, hi) in enumerate(ranges):
            center = (lo + hi) / 2
            dist = abs(value - center)
            if dist < best_dist:
                best_dist = dist
                best_idx = i
        return best_idx

    def _build_grid_from_cells(self, cells, table_bbox):
        """Fallback: dựng lưới bằng clustering x_center của cell có chữ.

        Logic cũ — chỉ dùng khi không có metadata row/col của detector.
        """
        valid_cells = [c for c in cells if 'bbox' in c and len(c['bbox']) == 4]
        if not valid_cells:
            return {"columns": 0, "rows": [], "col_widths": [], "row_heights": []}

        valid_cells.sort(key=lambda c: c['bbox'][1])

        rows = []
        current_row = [valid_cells[0]]
        for cell in valid_cells[1:]:
            cy = (cell['bbox'][1] + cell['bbox'][3]) / 2
            row_ymin = min(c['bbox'][1] for c in current_row)
            row_ymax = max(c['bbox'][3] for c in current_row)
            if row_ymin <= cy <= row_ymax:
                current_row.append(cell)
            else:
                rows.append(current_row)
                current_row = [cell]
        if current_row:
            rows.append(current_row)

        row_y_min = [min(c['bbox'][1] for c in r) for r in rows]
        row_y_max = [max(c['bbox'][3] for c in r) for r in rows]

        row_boundaries = [table_bbox.get('y1', 0)]
        for i in range(len(rows) - 1):
            boundary = (row_y_max[i] + row_y_min[i + 1]) / 2
            row_boundaries.append(boundary)
        row_boundaries.append(table_bbox.get('y2', 0))

        row_heights = [max(0, row_boundaries[i + 1] - row_boundaries[i])
                       for i in range(len(rows))]

        x_centers = []
        for r in rows:
            for c in r:
                x_centers.append((c['bbox'][0] + c['bbox'][2]) / 2)
        x_centers.sort()
        columns_x = []
        if x_centers:
            thresh = 20
            curr_col = [x_centers[0]]
            for xc in x_centers[1:]:
                if xc - sum(curr_col) / len(curr_col) < thresh:
                    curr_col.append(xc)
                else:
                    columns_x.append(sum(curr_col) / len(curr_col))
                    curr_col = [xc]
            if curr_col:
                columns_x.append(sum(curr_col) / len(curr_col))

        num_cols = len(columns_x)
        col_x_min = [float('inf')] * num_cols
        col_x_max = [0] * num_cols
        for row in rows:
            for cell in row:
                cx = (cell['bbox'][0] + cell['bbox'][2]) / 2
                col_idx = min(range(num_cols), key=lambda i: abs(cx - columns_x[i]))
                col_x_min[col_idx] = min(col_x_min[col_idx], cell['bbox'][0])
                col_x_max[col_idx] = max(col_x_max[col_idx], cell['bbox'][2])

        for i in range(num_cols):
            if col_x_min[i] == float('inf'):
                col_x_min[i] = columns_x[i] - 10
                col_x_max[i] = columns_x[i] + 10

        col_boundaries = [table_bbox.get('x1', 0)]
        for i in range(num_cols - 1):
            boundary = (col_x_max[i] + col_x_min[i + 1]) / 2
            col_boundaries.append(boundary)
        col_boundaries.append(table_bbox.get('x2', 0))
        col_widths = [max(0, col_boundaries[i + 1] - col_boundaries[i])
                      for i in range(num_cols)]

        grid = []
        for row in rows:
            row_texts = [""] * num_cols
            for cell in row:
                cx = (cell['bbox'][0] + cell['bbox'][2]) / 2
                col_idx = min(range(num_cols), key=lambda i: abs(cx - columns_x[i]))
                if row_texts[col_idx]:
                    row_texts[col_idx] += " " + cell.get('ocr_text', '')
                else:
                    row_texts[col_idx] = cell.get('ocr_text', '')
            grid.append([t.strip() for t in row_texts])

        return {
            "columns": num_cols,
            "rows": grid,
            "col_widths": col_widths,
            "row_heights": row_heights,
        }

    def _lock_table_layout(self, table, left_inch, top_inch, width_inch,
                           height_inch=None, height_rule='exact'):
        """Khóa vị trí/kích thước bảng nổi.

        height_rule:
            - 'exact'   : khóa cứng chiều cao bằng `height_inch` (Word cắt nội dung vượt).
            - 'atLeast' : dùng `height_inch` làm chiều cao tối thiểu, cho phép bảng
                          tự nở thêm khi nội dung vượt (phù hợp cho khối Notes vì OCR
                          có thể bể dòng nhiều hơn so với layout gốc).
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches
        tblPr = table._element.tblPr
        
        # Cho phép Word tự động nới rộng cột nếu text quá dài (AutoFit)
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'autofit')
        tblPr.append(tblLayout)
        
        # Cho phép đè lấp
        tblOverlap = OxmlElement('w:tblOverlap')
        tblOverlap.set(qn('w:val'), 'overlap')
        tblPr.append(tblOverlap)
        
        # Table position (WML: floating table)
        tblpPr = OxmlElement('w:tblpPr')
        tblpPr.set(qn('w:leftFromText'), '0')
        tblpPr.set(qn('w:rightFromText'), '0')
        tblpPr.set(qn('w:topFromText'), '0')
        tblpPr.set(qn('w:bottomFromText'), '0')
        tblpPr.set(qn('w:vertAnchor'), 'page')
        tblpPr.set(qn('w:horzAnchor'), 'page')
        tblpPr.set(qn('w:tblpX'), str(int(left_inch * 1440)))
        tblpPr.set(qn('w:tblpY'), str(int(top_inch * 1440)))
        tblPr.append(tblpPr)
        
        # Đặt lại margin của Table về 0
        tblCellMar = OxmlElement('w:tblCellMar')
        for m in ['top', 'left', 'bottom', 'right']:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), '0')
            node.set(qn('w:type'), 'dxa')
            tblCellMar.append(node)
        tblPr.append(tblCellMar)
        
        # Đặt Width tổng
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(int(width_inch * 1440)))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        
        # Ép chiều cao (nếu được truyền vào). Với 'atLeast' bảng có thể nở thêm
        # khi nội dung dài hơn vùng layout phát hiện được.
        if height_inch is not None:
            for row in table.rows:
                row.height = Inches(height_inch)
                trPr = row._tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), str(int(height_inch * 1440)))
                trHeight.set(qn('w:hRule'), height_rule)
                trPr.append(trHeight)

    def _export_to_docx(self, doc_struct):
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import qn, nsdecls

        doc = Document()
        
        width_px = doc_struct["document"]["width"]
        height_px = doc_struct["document"]["height"]
        image_name = doc_struct["document"]["image_name"]
        
        # Dùng tỷ lệ 100 pixels = 1 inch
        scale = 100.0
        page_width_inch = width_px / scale
        page_height_inch = height_px / scale
        
        section = doc.sections[0]
        section.page_width = Inches(page_width_inch)
        section.page_height = Inches(page_height_inch)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)
        
        # (Đã loại bỏ chức năng chèn ảnh nền gốc theo yêu cầu)

        # Thêm các element dưới dạng bảng nổi (floating table)
        for element in doc_struct["elements"]:
            etype = element["type"]
            bbox = element.get("bbox", [0, 0, 0, 0])
            
            # Thêm padding để che hết nét chữ cũ
            pad = 5
            x1 = max(0, bbox[0] - pad)
            y1 = max(0, bbox[1] - pad)
            w = bbox[2] - bbox[0] + pad * 2
            h = bbox[3] - bbox[1] + pad * 2
            
            left_inch = x1 / scale
            top_inch = y1 / scale
            width_inch = w / scale
            height_inch = h / scale
            
            if etype == "table":
                structure = element.get("structure", {})
                rows = structure.get("rows", [])
                num_cols = structure.get("columns", 0)
                col_widths = structure.get("col_widths", [])
                row_heights = structure.get("row_heights", [])
                
                if not rows or num_cols == 0:
                    continue
                    
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'
                
                # Cài đặt cột và chiều rộng
                for i, col in enumerate(table.columns):
                    if i < len(col_widths):
                        col_inch = (col_widths[i] + (pad*2/num_cols)) / scale
                        col.width = Inches(col_inch)
                        for cell in col.cells:
                            tcW = OxmlElement('w:tcW')
                            tcW.set(qn('w:w'), str(int(col_inch * 1440)))
                            tcW.set(qn('w:type'), 'dxa')
                            cell._tc.get_or_add_tcPr().append(tcW)
                            
                            vAlign = OxmlElement('w:vAlign')
                            vAlign.set(qn('w:val'), 'center')
                            cell._tc.get_or_add_tcPr().append(vAlign)

                # Ép tọa độ tuyệt đối và khóa cứng Layout bằng Helper
                self._lock_table_layout(table, left_inch, top_inch, width_inch)
                
                for i, row_data in enumerate(rows):
                    row = table.rows[i]
                    if i < len(row_heights):
                        row_inch = (row_heights[i] + (pad*2/len(rows))) / scale
                        row.height = Inches(row_inch)
                        trPr = row._tr.get_or_add_trPr()
                        trHeight = OxmlElement('w:trHeight')
                        trHeight.set(qn('w:val'), str(int(row_inch * 1440)))
                        trHeight.set(qn('w:hRule'), 'exact')
                        trPr.append(trHeight)

                    for j, cell_text in enumerate(row_data):
                        if j < len(table.columns):
                            cell = table.cell(i, j)
                            cell.text = cell_text
                            
                            # Background trắng
                            shading_elm = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))
                            cell._tc.get_or_add_tcPr().append(shading_elm)
                            
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.space_after = Pt(0)
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.size = Pt(8)
                                    run.font.name = 'Times New Roman'
                                    
            elif etype in ["notes", "metadata"]:
                raw_items = element.get("items", []) if etype == "notes" else [element.get("text", "")]
                items = [it for it in raw_items if it and it.strip()]
                if not items:
                    continue

                # Notes cần nhiều chỗ hơn vì OCR thường bể dòng khác với layout gốc.
                # Nới rộng padding ngang + dùng atLeast để bảng tự cao thêm nếu cần
                # → tránh tình trạng text bị cắt cụt như khi dùng hRule='exact'.
                note_pad = 12
                nx1 = max(0, bbox[0] - note_pad)
                ny1 = max(0, bbox[1] - note_pad)
                nw = bbox[2] - bbox[0] + note_pad * 2
                nh = bbox[3] - bbox[1] + note_pad * 2
                n_left = nx1 / scale
                n_top = ny1 / scale
                n_width = nw / scale
                n_height = nh / scale

                table = doc.add_table(rows=1, cols=1)
                self._lock_table_layout(table, n_left, n_top, n_width, n_height,
                                        height_rule='atLeast')

                tblPr = table._element.tblPr
                tblBorders = OxmlElement('w:tblBorders')
                for b_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{b_name}')
                    border.set(qn('w:val'), 'none')
                    tblBorders.append(border)
                tblPr.append(tblBorders)

                cell = table.cell(0, 0)
                shading_elm = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)

                cell.text = ""
                first_para = cell.paragraphs[0]
                for idx, item in enumerate(items):
                    if idx == 0:
                        p = first_para
                        p.text = item
                    else:
                        p = cell.add_paragraph(item)
                    p.paragraph_format.space_after = Pt(2)
                    for run in p.runs:
                        run.font.size = Pt(8)
                        run.font.name = 'Times New Roman'
                        if etype == "metadata":
                            run.font.color.rgb = RGBColor(100, 100, 100)
                            run.font.italic = True
                                
            elif etype == "drawing":
                img_path = element.get("image_path")
                if not img_path or not Path(img_path).exists():
                    continue
                    
                table = doc.add_table(rows=1, cols=1)
                
                # Ép tọa độ tuyệt đối và khóa cứng Layout bằng Helper
                self._lock_table_layout(table, left_inch, top_inch, width_inch, height_inch)
                
                tblPr = table._element.tblPr
                # Bỏ viền
                tblBorders = OxmlElement('w:tblBorders')
                for b_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{b_name}')
                    border.set(qn('w:val'), 'none')
                    tblBorders.append(border)
                tblPr.append(tblBorders)
                
                cell = table.cell(0, 0)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                # Chèn hình với chiều rộng bằng bbox
                try:
                    run.add_picture(str(img_path), width=Inches(width_inch))
                except:
                    p.add_run("[Lỗi load ảnh drawing]")
                                
        docx_path = self.output_dir / f"{Path(doc_struct['document']['image_name']).stem}_reconstructed.docx"
        doc.save(str(docx_path))
        print(f"Generated Reconstructed DOCX: {docx_path}")
